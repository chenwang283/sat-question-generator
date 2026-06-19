#!/usr/bin/env python3
"""
Build a .docx for ONE verified question-set markdown file — the 10 questions for a single
(skill x difficulty) — in the fixed layout from references/output-format.md: each item an
identical block, with the Math `spec` blocks stripped. The pipeline runs this once per
difficulty, producing three separate docs (easy / medium / hard) for a run.

(More than one file may be passed; they are grouped by skill, then difficulty, into one
doc. The pipeline does not do this — it builds one difficulty per doc.)

Usage (one doc per difficulty):
    python scripts/build_docx.py --out final/linear-equations-in-one-variable-easy.docx \
        --title "Linear equations in one variable - Easy"   final/linear-equations-in-one-variable__easy.md
    python scripts/build_docx.py --out final/linear-equations-in-one-variable-medium.docx \
        --title "Linear equations in one variable - Medium" final/linear-equations-in-one-variable__medium.md
    python scripts/build_docx.py --out final/linear-equations-in-one-variable-hard.docx \
        --title "Linear equations in one variable - Hard"   final/linear-equations-in-one-variable__hard.md

Requires python-docx:  pip install python-docx --break-system-packages
"""
import argparse, json, re, sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx --break-system-packages")
    sys.exit(2)

ROOT = Path(__file__).resolve().parent.parent
TAX = json.loads((ROOT / "references" / "taxonomy.json").read_text())

META_RE = re.compile(r"set-meta:\s*(.*)")
ITEM_RE = re.compile(r"^##\s+Item\s+\d+\b", re.MULTILINE)
# Multi-line fields capture up to the next bold label, the `---` separator, or end of
# block — so the correct-answer explanation no longer swallows the `Why not X:` lines.
def _tail(label):
    return re.compile(rf"^\*\*{label}:\*\*\s*(.*?)(?=\n\*\*|\n---|\Z)", re.MULTILINE | re.DOTALL)

FIELD_RE = {
    "passage":     re.compile(r"^\*\*Passage:\*\*\s*(.*)", re.MULTILINE),
    "stem":        re.compile(r"^\*\*Stem:\*\*\s*(.*)", re.MULTILINE),
    "A":           re.compile(r"^\*\*A\.\*\*\s*(.*)", re.MULTILINE),
    "B":           re.compile(r"^\*\*B\.\*\*\s*(.*)", re.MULTILINE),
    "C":           re.compile(r"^\*\*C\.\*\*\s*(.*)", re.MULTILINE),
    "D":           re.compile(r"^\*\*D\.\*\*\s*(.*)", re.MULTILINE),
    "correct":     re.compile(r"^\*\*Correct:\*\*\s*([A-D])", re.MULTILINE),
    "explanation": _tail("Explanation"),
    "why_not_A":   _tail("Why not A"),
    "why_not_B":   _tail("Why not B"),
    "why_not_C":   _tail("Why not C"),
    "why_not_D":   _tail("Why not D"),
}
# Fields that live in the tail of an item (after the options) and must be read from the
# spec-stripped text so a Math item's trailing `spec` JSON never bleeds into them.
TAIL_KEYS = {"explanation", "why_not_A", "why_not_B", "why_not_C", "why_not_D"}
DIFF_ORDER = {"easy": 0, "medium": 1, "hard": 2}


def label_for(subject, topic, skill):
    try:
        return TAX[subject]["topics"][topic]["skills"][skill]
    except KeyError:
        return skill.replace("-", " ").capitalize()


def parse_meta(md):
    m = META_RE.search(md)
    meta = {}
    if m:
        for kv in m.group(1).split(";"):
            if "=" in kv:
                k, v = kv.split("=", 1)
                meta[k.strip()] = v.strip()
    return meta


def parse_items(md):
    starts = [m.start() for m in ITEM_RE.finditer(md)]
    items = []
    for i, start in enumerate(starts):
        end = starts[i + 1] if i + 1 < len(starts) else len(md)
        block = md[start:end]
        before_spec = re.split(r"```spec", block)[0]   # keep the spec JSON out of the tail fields
        item = {}
        for key, rx in FIELD_RE.items():
            src = before_spec if key in TAIL_KEYS else block
            mm = rx.search(src)
            if mm:
                item[key] = mm.group(1).strip()
        if "stem" in item:
            items.append(item)
    return items


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("inputs", nargs="+", help="question-set markdown files")
    ap.add_argument("--out", required=True)
    ap.add_argument("--title", default="SAT Practice Set")
    ap.add_argument("--answer-key-section", action="store_true",
                    help="append a compact answer key after all questions")
    args = ap.parse_args()

    grouped, labels = {}, {}
    for path in args.inputs:
        md = Path(path).read_text()
        meta = parse_meta(md)
        subject = meta.get("subject", "")
        topic = meta.get("topic", "")
        skill = meta.get("skill", Path(path).stem)
        diff = meta.get("difficulty", "unspecified")
        key = (subject, topic, skill)
        labels[key] = label_for(subject, topic, skill)
        grouped.setdefault(key, {})[diff] = parse_items(md)

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    title = doc.add_paragraph()
    r = title.add_run(args.title)
    r.bold = True
    r.font.size = Pt(18)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    answer_key = []
    for key in sorted(grouped, key=lambda k: (k[0], k[1], k[2])):
        doc.add_heading(labels[key], level=1)
        diffs = grouped[key]
        for diff in sorted(diffs, key=lambda d: DIFF_ORDER.get(d, 99)):
            doc.add_heading(diff.capitalize(), level=2)
            for n, item in enumerate(diffs[diff], 1):
                qp = doc.add_paragraph()
                qp.add_run(f"Question {n}").bold = True
                if item.get("passage"):
                    doc.add_paragraph(item["passage"])
                doc.add_paragraph(item.get("stem", ""))
                for letter in ("A", "B", "C", "D"):
                    if letter in item:
                        doc.add_paragraph(f"{letter}. {item[letter]}")
                ans = doc.add_paragraph()
                ans.add_run("Answer: ").bold = True
                ans.add_run(item.get("correct", ""))
                exp = doc.add_paragraph()
                exp.add_run("Explanation: ").bold = True
                exp.add_run(item.get("explanation", ""))
                for letter in ("A", "B", "C", "D"):
                    why = item.get(f"why_not_{letter}")
                    if why:
                        wp = doc.add_paragraph()
                        wp.add_run(f"Why not {letter}: ").bold = True
                        wp.add_run(why)
                doc.add_paragraph("")  # spacer between questions
                answer_key.append((labels[key], diff, n, item.get("correct", "")))

    if args.answer_key_section and answer_key:
        doc.add_page_break()
        doc.add_heading("Answer Key", level=1)
        cur = None
        for sub_label, diff, n, correct in answer_key:
            head = f"{sub_label} — {diff.capitalize()}"
            if head != cur:
                doc.add_heading(head, level=2)
                cur = head
            doc.add_paragraph(f"{n}. {correct}")

    Path(args.out).parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.out)
    print(f"Wrote {args.out}: {len(grouped)} skill(s), {len(answer_key)} item(s).")


if __name__ == "__main__":
    main()
